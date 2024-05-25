import os
from google.cloud import documentai_v1beta3 as documentai
from google.oauth2 import service_account
from docx import Document
from fpdf import FPDF
from dotenv import load_dotenv

load_dotenv()

class DocumentExtractor:
    def __init__(self):
        credentials_path = os.getenv("GOOGLE_APPLICATION_CREDENTIALS")
        self.credentials = service_account.Credentials.from_service_account_file(credentials_path)
        self.client = documentai.DocumentProcessorServiceClient(credentials=self.credentials)
        self.location = os.getenv("LOCATION")
        self.processor_id = os.getenv("PROCESSOR_ID")
        self.project_id = os.getenv("PROJECT_ID")
        self.endpoint = os.getenv("ENDPOINT")
        self.information = None
        

    def convert_word_to_pdf(self, docx_path, pdf_path):
        document = Document(docx_path)
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("Arial", size=12)

        for paragraph in document.paragraphs:
            pdf.multi_cell(0, 10, paragraph.text)
        
        pdf.output(pdf_path)

    def extract_text_from_pdf(self, pdf_path):
        try:
            mime_type = 'application/pdf'
            name = self.client.processor_path(self.project_id, self.location, self.processor_id)
            with open(pdf_path, "rb") as image:
                image_content = image.read()
            
            raw_document = documentai.types.RawDocument(
                content=image_content, mime_type=mime_type)
            
            request = {"name": name, "raw_document": raw_document}
            response = self.client.process_document(request=request)
            document = response.document
            self.information = document
            
            return document.text
        except Exception as e:
            print(e)
            return None
"""
if __name__ == '__main__':
    
    extractor = DocumentExtractor()

    pdf_file_path = 'uploads/pdf/Auto.pdf'

    text = extractor.extract_text_from_pdf(pdf_file_path)
    print(text)"""