import google.generativeai as genai
from PIL import Image
import os
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from io import BytesIO
from google.cloud import documentai_v1beta3 as documentai
from dotenv import load_dotenv


load_dotenv()

class geminiApi:
    def __init__(self):
        self.api_key = os.getenv("GEMINI_API_KEY")
        self.model = None

    def configure(self):
        genai.configure(api_key=self.api_key)
        self.model = genai.GenerativeModel(model_name="gemini-pro-vision")
        return self.model

    def configure_2(self):
        genai.configure(api_key=self.api_key)
        self.model = genai.GenerativeModel(model_name="gemini-pro")
        return self.model
    
    def extract_text_from_pdf(self, pdf_path):
        try:
            client = documentai.DocumentUnderstandingServiceClient()
            with open(pdf_path, "rb") as image:
                image_content = image.read()

            name = os.path.basename(pdf_path)
            mime_type = "application/pdf"
            parent = f"projects/{self.project_id}/locations/{self.location}/processors/{self.processor_id}"
            raw_document = {"content": image_content, "mime_type": mime_type}

            request = {"name": name, "raw_document": raw_document}
            response = client.process_document(request=request, parent=parent)
            document = response.document
            text = document.text

            return text
        except Exception as e:
            print(e)
            return None
    def generate_user_story(self, file_path, description, file_type):
        if file_type == 'image':
            imagen = Image.open(file_path)
            prompt = f"{description}, quiero que formules dos historias de usuario con el formato de como, quiero, para (no debe especificar la sintaxis)"
            response = self.model.generate_content([prompt, imagen])
        elif file_type == 'pdf':
            text = self.extract_text_from_pdf(file_path)
            prompt = f"{description} {text}, quiero que formules dos historias de usuario con el formato de como, quiero, para (no debe especificar la sintaxis)"
            response = self.model.generate_content(prompt)
        elif file_type == 'docx':
            text = self.extract_text_from_docx(file_path)
            prompt = f"{description} {text}, quiero que formules dos historias de usuario con el formato de como, quiero, para (no debe especificar la sintaxis)"
            response = self.model.generate_content(prompt)
        else:
            return []

        historias_usuario = response.text.split('\n')
        return [historia.strip() for historia in historias_usuario if historia.strip()]
    
    def generate_user_story_info(self, info_document):
        prompt = info_document
        estructura = 'dame solo una historia de usuario con el formato de como, quiero, para'
        full_prompt = f'{estructura} {prompt}'
        response = self.model.generate_content(full_prompt)
        return response.text
    
     #Metodos de obtención de información mediante texto con respecto a otras funcionalidades 
    def generate_requirements(self, info):
        prompt = info
        estructura = 'Con esta estructura quiero 5 requisitos funcionales: [Sujeto]-[Acción]-[Valor]. y 5 requisitos no funcionales : [Condición]-[Sujeto]-[Acción]-[Objeto]-[Restricción] Con respecto a estas Historias de usuario,hazlo con un lenguaje natural usando articulos y pronombres en una frase, y ademas agregale despues del sujeto un "debe ser"  : '
        full_prompt = f'{estructura} {prompt}'
        response = self.model.generate_content(full_prompt)
        requisitos_usuario = response.text.split('\n')
        return [requisitos.strip() for requisitos in requisitos_usuario if requisitos.strip()]

    def generate_classification_prioritization(self, text_R):
        estructura = 'Prioriza los siguientes requisitos según su importancia, enumerándolos del más importante al menos importante:'
        full_prompt = f'{estructura} {text_R}'
        response = self.model.generate_content(full_prompt)
        prioritization = response.text
        #print("Priorización de requisitos generada:")
        #print(prioritization)

        # Parsear la respuesta para obtener la lista de requisitos priorizados
        prioritized_requirements = [req.strip() for req in prioritization.split('\n') if req.strip()]

        return prioritized_requirements
    
    def extract_text_from_pdf(self, pdf_path):
        try:
            mime_type = 'application/pdf'
            name = self.client.processor_path(self.project_id, self.location, self.processor_id)
            with open(pdf_path, "rb") as image:
                image_content = image.read()

            raw_document = documentai.types.RawDocument(
                content=image_content, mime_type=mime_type)

            request = {"name": name, "raw_document": raw_document}
            response = self.client.process_document(request=request)
            document = response.document
            text = document.text

            self.configure_2()
            return text
        except Exception as e:
            print(e)
            return None

    def extract_text_from_docx(self, docx_path):
        try:
            document = Document(docx_path)
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)

            self.configure_2()
            return text
        except Exception as e:
            print(e)
            return None






def generate_word_document(historias_usuario):
    doc = Document()
    doc.add_heading('Historias de Usuario', 0)
    for historia in historias_usuario:
        doc.add_paragraph(historia)
    byte_io = BytesIO()
    doc.save(byte_io)
    byte_io.seek(0)
    return byte_io.getvalue()

def generate_pdf_document(historias_usuario):
    byte_io = BytesIO()
    c = canvas.Canvas(byte_io, pagesize=letter)
    width, height = letter
    c.drawString(100, height - 40, "Historias de Usuario")
    y = height - 60
    for historia in historias_usuario:
        c.drawString(100, y, historia)
        y -= 20
        if y < 40:
            c.showPage()
            y = height - 40
    c.save()
    byte_io.seek(0)
    return byte_io.getvalue()

def generate_txt_document(historias_usuario):
    content = "\n".join(historias_usuario)
    return content.encode('utf-8')