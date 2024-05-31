import os
from google.cloud import documentai_v1beta3 as documentai
from google.oauth2 import service_account
from docx import Document
from fpdf import FPDF
from dotenv import load_dotenv
from .gemini import geminiApi

load_dotenv()

class DocumentExtractor:
    def __init__(self):
        credentials_path = os.getenv("GOOGLE_APPLICATION_CREDENTIALS")
        self.credentials = service_account.Credentials.from_service_account_file(credentials_path)
        self.client = documentai.DocumentProcessorServiceClient(credentials=self.credentials)
        self.location = os.getenv("LOCATION")
        self.processor_id = os.getenv("PROCESSOR_ID")
        self.project_id = os.getenv("PROJECT_ID")
        self.endpoint = os.getenv("ENDPOINT")
        self.geminiApi = geminiApi()
        

    def convert_word_to_pdf(self, docx_path, pdf_path):
        document = Document(docx_path)
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("Arial", size=12)

        for paragraph in document.paragraphs:
            pdf.multi_cell(0, 10, paragraph.text)
        
        pdf.output(pdf_path)

    def extract_text_from_pdf(self, pdf_path):
        try:
            mime_type = 'application/pdf'
            name = self.client.processor_path(self.project_id, self.location, self.processor_id)
            with open(pdf_path, "rb") as image:
                image_content = image.read()
            
            raw_document = documentai.types.RawDocument(
                content=image_content, mime_type=mime_type)
            
            request = {"name": name, "raw_document": raw_document}
            response = self.client.process_document(request=request)
            document = response.document
            text = document.text

            self.geminiApi.configure_2()
            return text
        except Exception as e:
            print(e)
            return None
        
    def extract_text_from_docx(self, docx_path):
        try:
            document = Document(docx_path)
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)

            self.geminiApi.configure_2()
            return text
        except Exception as e:
            print(e)
            return None
        
    def extract_text_from_txt(self, txt_path):
        try:
            with open(txt_path, 'r', encoding='utf-8') as file:
                text = file.read()

            self.geminiApi.configure_2()
            return text
        except Exception as e:
            print(e)
            return None
        
    def extract_prioritized_requirements(self, file_path, description):
        file_extension = os.path.splitext(file_path)[1].lower()
        if file_extension == '.pdf':
            text = self.extract_text_from_pdf(file_path)
        elif file_extension == '.docx':
            text = self.extract_text_from_docx(file_path)
        elif file_extension == '.txt':
            text = self.extract_text_from_txt(file_path)
        else:
            return "Formato de archivo no soportado."

        if not text:
            return "No se pudo extraer texto del archivo."

        # Incluir la descripción del usuario en la solicitud a la API de Gemini
        full_text = f"Descripción del usuario: {description}\n\n{text}"
        prioritized_requirements = self.geminiApi.generate_classification_prioritization(full_text)
        return prioritized_requirements
    
    def clasification_requirements(self, file_path, description):
        file_extension = os.path.splitext(file_path)[1].lower()
        if file_extension == '.pdf':
            text = self.extract_text_from_pdf(file_path)
        elif file_extension == '.docx':
            text = self.extract_text_from_docx(file_path)
        elif file_extension == '.txt':
            text = self.extract_text_from_txt(file_path)
        else:
            return "Formato de archivo no soportado."

        if not text:
            return "No se pudo extraer texto del archivo."

        # Incluir la descripción del usuario en la solicitud a la API de Gemini
        full_text = f"Descripción del usuario: {description}\n\n{text}"
        clasification_requirements = self.geminiApi.generate_requirements(full_text)
        return clasification_requirements
"""
extract = DocumentExtractor()

docx_path = 'uploads/pdf/Uno.docx'
pdf_path = 'uploads/pdf/Uno_converted.pdf'

# Convertir DOCX a PDF
extract.convert_word_to_pdf(docx_path, pdf_path)

# Procesar el archivo PDF convertido
extract.extract_text_from_pdf(pdf_path)
print(extract.geminiApi.text)
"""