from flask import Blueprint, render_template, current_app, flash, send_file, request, redirect, url_for, session
from form import MyFormDocument
from werkzeug.utils import secure_filename
import os
from io import BytesIO
from classe.document import DocumentExtractor
from classe.gemini import geminiApi
import atexit

home_bp = Blueprint("home", __name__, template_folder="templates")

@home_bp.route('/')
def home():
    return render_template('home.html')

@home_bp.route('/clasificacion', methods=['GET', 'POST'])
def clasificacion():
    form = MyFormDocument()
    if form.validate_on_submit():
        description = form.description.data
        file = form.file.data
        if not description:
            flash('Por favor, proporcione una descripción.', 'error')
        elif not file:
            flash('Por favor, suba un documento.', 'error')
        else:
            filename = secure_filename(file.filename)
            file_path = os.path.join(current_app.config['UPLOADED_DOCUMENTS_DEST'], filename)
            
            os.makedirs(os.path.dirname(file_path), exist_ok=True)
            file.save(file_path)
            
            extractor = DocumentExtractor()
            clasifica_requirements = extractor.clasification_requirements(file_path, description)
            
            if clasifica_requirements:
                flash('Requisitos clasificados correctamente.', 'success')
                session['textos'] = clasifica_requirements  # Almacenar en sesión
                return render_template('clasificacion.html', form=form, textos=clasifica_requirements)
            else:
                flash('Error al procesar el documento.', 'error')
    return render_template('clasificacion.html', form=form)


@home_bp.route('/priorizacion', methods=['GET', 'POST'])
def priorizacion():
    form = MyFormDocument()
    if form.validate_on_submit():
        description = form.description.data
        file = form.file.data
        if not description:
            flash('Por favor, proporcione una descripción.', 'error')
        elif not file:
            flash('Por favor, suba un documento.', 'error')
        else:
            filename = secure_filename(file.filename)
            file_path = os.path.join(current_app.config['UPLOADED_DOCUMENTS_DEST'], filename)
            
            os.makedirs(os.path.dirname(file_path), exist_ok=True)
            file.save(file_path)
            
            extractor = DocumentExtractor()
            prioritized_requirements = extractor.extract_prioritized_requirements(file_path, description)
            
            if prioritized_requirements:
                flash('Requisitos priorizados obtenidos correctamente.', 'success')
                session['textos'] = prioritized_requirements  # Almacenar en sesión
                return render_template('priorizacion.html', form=form, textos=prioritized_requirements)
            else:
                flash('Error al procesar el documento.', 'error')

    return render_template('priorizacion.html', form=form)

@home_bp.route('/HU_imagen', methods=['GET', 'POST'])
def HU_imagen():
    form = MyFormDocument()
    if form.validate_on_submit():
        description = form.description.data
        file = form.file.data
        if not description:
            flash('Por favor, proporcione una descripción.', 'error')
        elif not file:
            flash('Por favor, suba un archivo.', 'error')
        else:
            filename = secure_filename(file.filename)
            file_path = os.path.join(current_app.config['GENERATED_UPLOADS_FOLDER'], filename)
            
            os.makedirs(os.path.dirname(file_path), exist_ok=True)
            file.save(file_path)

            file_type = None
            if file.filename.lower().endswith(('.png', '.jpg', '.jpeg')):
                file_type = 'image'
            elif file.filename.lower().endswith('.pdf'):
                file_type = 'pdf'
            elif file.filename.lower().endswith('.docx'):
                file_type = 'docx'
            
            if file_type:
                api = geminiApi()
                api.configure()
                textos = api.generate_user_story(file_path, description, file_type)
                
                session['textos'] = textos  # Almacenar en sesión
                
                return render_template('HU_imagen.html', form=form, textos=textos)
            else:
                flash('Formato de archivo no soportado.', 'error')

    return render_template('HU_imagen.html', form=form, textos=None)

@home_bp.route('/download_document/<format>', methods=['GET'])
def download_document(format):
    textos = session.get('textos', [])
    
    if not textos:
        flash('No hay historias de usuario disponibles para descargar.', 'error')
        return redirect(url_for('home.priorizacion'))

    if format == 'word':
        document_content = generate_word_document(textos)
        filename = 'textos.docx'
        mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    elif format == 'pdf':
        document_content = generate_pdf_document(textos)
        filename = 'textos.pdf'
        mimetype = 'application/pdf'
    elif format == 'txt':
        document_content = generate_txt_document(textos)
        filename = 'textos.txt'
        mimetype = 'text/plain'
    else:
        return 'Formato no admitido', 400

    output_folder = current_app.config['GENERATED_UPLOADS_FOLDER']
    os.makedirs(output_folder, exist_ok=True)

    file_path = os.path.join(output_folder, filename)
    with open(file_path, 'wb') as f:
        f.write(document_content)

    # Guarda el archivo generado en la lista global
    generated_files.append(file_path)

    return send_file(file_path, as_attachment=True, download_name=filename, mimetype=mimetype)

def generate_word_document(textos):
    from docx import Document
    document = Document()
    for historia in textos:
        document.add_paragraph(historia)
    byte_io = BytesIO()
    document.save(byte_io)
    byte_io.seek(0)  # Reset the stream position to the beginning
    return byte_io.getvalue()

def generate_pdf_document(textos):
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    byte_io = BytesIO()
    c = canvas.Canvas(byte_io, pagesize=letter)
    y = 750
    for historia in textos:
        c.drawString(40, y, historia)
        y -= 20
    c.showPage()
    c.save()
    byte_io.seek(0)  # Reset the stream position to the beginning
    return byte_io.getvalue()

def generate_txt_document(textos):
    content = "\n".join(textos)
    return content.encode('utf-8')

# Variable global para almacenar archivos generados
generated_files = []

# Función para limpiar archivos generados
def cleanup_files():
    for file_path in generated_files:
        try:
            os.remove(file_path)
        except OSError:
            pass

# Registrar la función de limpieza para ejecutarse al cerrar la aplicación
import atexit
atexit.register(cleanup_files)